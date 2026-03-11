"""
Module 4 — CAM Generator
Generates a professional Credit Appraisal Memo as .docx using python-docx.
7 sections with embedded matplotlib charts.
"""

import io
import uuid
from datetime import datetime
from pathlib import Path


def generate_cam_document(session_id: str, session: dict) -> str:
    """Generate CAM .docx and return the file path."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[CAM] python-docx not installed. Run: pip install python-docx")
        return ""

    doc = Document()

    # Set up styles
    _setup_styles(doc)

    jury = session.get("jury_result", {})
    five_cs = jury.get("five_cs", session.get("five_cs", {}))
    verdict = jury.get("verdict", {})
    entity = session.get("entity", {})
    risk_flags = session.get("risk_flags", [])

    base_score = jury.get("base_score", 78.0)
    jury_score = jury.get("jury_score", 51.0)
    prosecution = jury.get("prosecution", {}).get("prosecution_findings", [])
    defence_findings = jury.get("defence", {}).get("defence_findings", [])

    # ─── Header ───────────────────────────────────────────────────────────────
    h = doc.add_heading("CREDIT APPRAISAL MEMO", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(h, "1B2B4B")

    sub = doc.add_paragraph("Intelli-Credit Automated Analysis System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x4A, 0x9D, 0xB5)

    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}   |   Session ID: {session_id[:8].upper()}")
    doc.add_paragraph("━" * 60)

    # ─── Section 1 — Borrower Profile ───────────────────────────────────────
    _section_heading(doc, "SECTION 1 — BORROWER PROFILE")
    _two_col(doc, "Company Name", entity.get("company_name", "Vardhaman Infra & Logistics Pvt. Ltd."))
    _two_col(doc, "CIN", entity.get("cin", "U45200MH2015PTC264831"))
    _two_col(doc, "GSTIN", entity.get("gstin", "27AABCV1234F1ZA"))
    _two_col(doc, "Sector / Industry", (entity.get("sector", "Infrastructure / Road Construction")).replace("_", " ").title())
    _two_col(doc, "Loan Amount Applied", f"₹{session.get('loan_amount_cr', 60)} Crore")
    _two_col(doc, "Loan Type", session.get("loan_type", "Term Loan"))
    _two_col(doc, "Date of Application", datetime.now().strftime("%d %B %Y"))

    # ─── Section 2 — Executive Summary ──────────────────────────────────────
    _section_heading(doc, "SECTION 2 — EXECUTIVE SUMMARY")
    rec = jury.get("jury_decision", "CONDITIONAL")
    p = doc.add_paragraph()
    p.add_run(f"RECOMMENDATION: {rec}").bold = True

    for factor in verdict.get("decisive_factors", [])[:3]:
        doc.add_paragraph(f"• {factor}")

    rationale = verdict.get("verdict_rationale", "Based on comprehensive analysis.")
    p = doc.add_paragraph(rationale)
    p.runs[0].italic = True

    # ─── Section 3 — Five Cs ─────────────────────────────────────────────────
    _section_heading(doc, "SECTION 3 — FIVE Cs ASSESSMENT")

    c_details = {
        "C1": ("CHARACTER", 0.25, "Promoter background, litigation status, DIN compliance"),
        "C2": ("CAPACITY", 0.30, "Revenue trend, EBITDA, DSCR, GST-bank match"),
        "C3": ("CAPITAL", 0.15, "Net worth, gearing ratio, equity buffer"),
        "C4": ("COLLATERAL", 0.20, "Asset type, estimated value, ROC charge status"),
        "C5": ("CONDITIONS", 0.10, "Sector outlook, regulatory risk, macro environment"),
    }

    for c_key, (c_name, weight, desc) in c_details.items():
        score = five_cs.get(c_key, 70)
        doc.add_heading(f"{c_key} {c_name} — Score: {score}/100  (Weight: {weight*100:.0f}%)", level=2)
        doc.add_paragraph(f"Key Inputs: {desc}")
        c_flags = [f for f in risk_flags if f.get("five_c_pillar") == c_key]
        if c_flags:
            for flag in c_flags[:3]:
                sev = flag.get("severity", "")
                p = doc.add_paragraph(f"  [{sev}] {flag.get('flag_type')}: {flag.get('evidence_snippet', '')[:150]}")

    # Embed Five Cs chart
    try:
        from modules.cam.chart_generator import create_five_cs_radar, create_gst_bank_chart, create_score_journey_chart
        radar_png = create_five_cs_radar(five_cs)
        if radar_png:
            buf = io.BytesIO(radar_png)
            doc.add_picture(buf, width=Inches(4))
    except Exception as e:
        print(f"[CAM] Chart failed: {e}")

    # ─── Section 4 — Research Intelligence ───────────────────────────────────
    _section_heading(doc, "SECTION 4 — RESEARCH INTELLIGENCE FINDINGS")
    research_flags = session.get("research_flags", [])
    high_med = [f for f in research_flags if f.get("severity") in ("HIGH", "MEDIUM")]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Finding"
    hdr[2].text = "Severity"
    hdr[3].text = "Date"

    for f in high_med[:10]:
        row = table.add_row().cells
        row[0].text = f.get("source_document", "")[:30]
        row[1].text = f.get("evidence_snippet", "")[:120]
        row[2].text = f.get("severity", "")
        row[3].text = str(f.get("created_at", ""))[:10]

    # ─── Section 5 — Jury Deliberation ───────────────────────────────────────
    _section_heading(doc, "SECTION 5 — AI AGENT JURY DELIBERATION")

    doc.add_paragraph(f"Base Model Score (XGBoost): {base_score}/100")
    doc.add_paragraph(f"Jury Consensus Score:       {jury_score}/100")
    doc.add_paragraph(f"Score Movement:             {jury.get('net_delta', 0)} points")

    doc.add_heading("Prosecution Findings (Risk Scrutiny Agent):", level=3)
    for pf in prosecution[:5]:
        doc.add_paragraph(f"  [{pf.get('finding_id')}] {pf.get('finding_text')} (Δ {pf.get('score_delta')} pts, {pf.get('five_c_pillar')})")

    doc.add_heading("Defence Findings (Credit Advocacy Agent):", level=3)
    for df in defence_findings[:5]:
        doc.add_paragraph(f"  [{df.get('finding_id')}] {df.get('finding_text')} (Δ +{df.get('score_delta')} pts, {df.get('five_c_pillar')})")

    doc.add_heading("Adjudication Rationale:", level=3)
    doc.add_paragraph(verdict.get("verdict_rationale", ""))

    # Score journey chart
    try:
        after_p = base_score + jury.get("prosecution_delta", 0)
        after_d = after_p + jury.get("defence_delta", 0)
        journey_png = create_score_journey_chart(base_score, after_p, after_d, jury_score)
        if journey_png:
            buf = io.BytesIO(journey_png)
            doc.add_picture(buf, width=Inches(6))
    except Exception as e:
        print(f"[CAM] Score journey chart failed: {e}")

    # ─── Section 6 — Final Decision ──────────────────────────────────────────
    _section_heading(doc, "SECTION 6 — FINAL DECISION")

    rec = jury.get("final_recommendation", "CONDITIONAL")
    amount = jury.get("recommended_loan_amount_cr", 35)
    rate = jury.get("recommended_interest_rate_pct", 12.5)
    tenor = jury.get("loan_tenor_months", 60)
    band = jury.get("confidence_band", {})

    _two_col(doc, "Decision", rec)
    _two_col(doc, "Recommended Amount", f"₹{amount} Crore  (Applied: ₹{session.get('loan_amount_cr', 60)} Cr)")
    _two_col(doc, "Confidence Band", f"₹{band.get('low', amount-5)}–{band.get('high', amount+5)} Crore")
    _two_col(doc, "Interest Rate", f"{rate}% p.a.")
    _two_col(doc, "Tenor", f"{tenor} months")

    doc.add_heading("Conditions Precedent:", level=3)
    for i, cond in enumerate(verdict.get("conditions", []), 1):
        doc.add_paragraph(f"{i}. {cond}")

    doc.add_paragraph("Next Steps:")
    doc.add_paragraph("1. Credit Officer to verify conditions with borrower")
    doc.add_paragraph("2. Legal team to review collateral documentation")
    doc.add_paragraph("3. Final sanction subject to Credit Committee ratification")

    # ─── Section 7 — Audit Trail ─────────────────────────────────────────────
    _section_heading(doc, "SECTION 7 — AUDIT TRAIL")
    _two_col(doc, "Session ID", session_id)
    _two_col(doc, "Analysis Timestamp", datetime.now().isoformat())
    _two_col(doc, "Credit Officer ID", session.get("qualitative", {}).get("officer_id", "N/A"))
    _two_col(doc, "AI Models Used", f"Claude claude-sonnet-4-20250514 / XGBoost {_get_xgboost_version()}")

    doc.add_paragraph("\nNote: This CAM was generated by Intelli-Credit AI system. All findings are traceable to source documents. Final credit decision requires authorised Credit Officer sign-off. AI recommendation is advisory only.", style="Normal")

    # Save
    out_dir = Path("cam_output")
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"CAM_{session_id}.docx"
    doc.save(str(out_path))
    print(f"[CAM] Generated: {out_path}")
    return str(out_path)


def _setup_styles(doc):
    from docx.shared import Pt
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _section_heading(doc, title: str):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc.add_paragraph("━" * 60)
    h = doc.add_heading(title, level=1)
    _set_heading_color(h, "1B2B4B")
    doc.add_paragraph("")


def _two_col(doc, key: str, value: str):
    p = doc.add_paragraph()
    r = p.add_run(f"{key:<25}")
    r.bold = True
    p.add_run(str(value))


def _set_heading_color(heading, hex_color: str):
    from docx.shared import RGBColor
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _get_xgboost_version() -> str:
    try:
        import xgboost as xgb
        return xgb.__version__
    except Exception:
        return "N/A"
